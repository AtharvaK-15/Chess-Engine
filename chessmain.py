import pygame as p
import chessEngine
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WIDTH = HEIGHT = 512
DIMENSION = 8
SQ_SIZE = HEIGHT // DIMENSION
MAX_FPS = 15
IMAGES = {}

'''
Initialize a global dictionary of images. This will be called exactly once in the main
'''

def loadImages():
    pieces = ['wP', 'wR', 'wN', 'wB', 'wQ', 'wK', 'bP', 'bR', 'bN', 'bB', 'bQ', 'bK']
    for piece in pieces:
        IMAGES[piece] = p.transform.scale(p.image.load("./images/" + piece + ".png"), (SQ_SIZE, SQ_SIZE))
    # Note: we can access an image by saying 'IMAGES['wp']'
        
'''
The main driver for our code. This will handle user input and updating the graphics
'''

def main():
    p.init()
    screen = p.display.set_mode((WIDTH, HEIGHT))
    clock = p.time.Clock()
    screen.fill(p.Color("white"))
    gs = chessEngine.GameState()
    validMoves = gs.getValidMoves()
    moveMade = False # flag variable for when a move is made
    doc = Document()
    print("White Side", end = " | ")
    print("Black Side")
    count = 0
    title = doc.add_heading("Notations", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    cell.text = "White Side"
    cell = table.cell(0, 1)
    cell.text = "Black Side"
    r=1
    c=0

    loadImages() # only do this once, before the while loop
    running = True
    sqSelected = () # no square is selected, keep track of the last click of the user (tuple: (row, col))
    playerClicks = [] # keep track of player clicks (two tuples: [(6, 4), (4, 4)])
    while running:
        for e in p.event.get():
            doc.save("notations.docx")
            if e.type == p.QUIT:
                running = False
            elif e.type == p.MOUSEBUTTONDOWN:
                location = p.mouse.get_pos()
                col = location[0]//SQ_SIZE
                row = location[1]//SQ_SIZE
                if sqSelected == (row, col):
                    sqSelected = ()
                    playerClicks = []
                else:
                    sqSelected = (row, col)
                    playerClicks.append(sqSelected) # append for both 1st and 2nd clicks
                if len(playerClicks) == 2:
                    move = chessEngine.Move(playerClicks[0], playerClicks[1], gs.board)
                    
                    for i in range(len(validMoves)):
                        if move == validMoves[i]:
                            gs.makeMove(validMoves[i])
                            moveMade = True
                            sqSelected = ()
                            playerClicks = []
                            table.add_row()
                            cell = table.cell(r, c)
                            check = 0                            
                            result = gs.inCheck()
                            if result:
                                check = 1

                            if(count%2 == 0):
                                if(check == 1):
                                        cell.text = str(count//2+1) + ". " + move.getChessNotation()+"+"
                                else:
                                        cell.text = str(count//2+1) + ". " + move.getChessNotation()
                                count+=1
                                c=1
                            else:
                                if(check == 1):
                                        cell.text = move.getChessNotation()+"+"
                                else:
                                    cell.text = move.getChessNotation()
                                count+=1
                                r+=1
                                c=0

                    if not moveMade:
                        playerClicks = [sqSelected]     
                
            elif e.type == p.KEYDOWN:
                if e.key == p.K_z: # undo when 'z' is pressed
                    gs.undoMove()
                    moveMade = True
                    for paragraph in cell.paragraphs:
                        paragraph.clear()
                    if(count%2 == 0):
                        r-=1
                        c=1
                        count-=1
                    else:
                        c=0
                        count-=1
        if moveMade:
            validMoves = gs.getValidMoves()
            moveMade = False

        drawGameState(screen, gs)
        clock.tick(MAX_FPS)
        p.display.flip()

def drawGameState(screen, gs):
    drawBoard(screen) # draw squares on the board
    drawPieces(screen, gs.board) # draw pieces on top of those squares

def drawBoard(screen):
    colors = [p.Color("white"), p.Color("gray")]
    for r in range(DIMENSION):
        for c in range(DIMENSION):
            color = colors[((r+c) % 2)]
            p.draw.rect(screen, color, p.Rect(c*SQ_SIZE, r*SQ_SIZE, SQ_SIZE, SQ_SIZE))

def drawPieces(screen, board):
    for r in range(DIMENSION):
        for c in range(DIMENSION):
            piece = board[r][c]
            if piece != "--": # not empty square
                screen.blit(IMAGES[piece], p.Rect(c*SQ_SIZE, r*SQ_SIZE, SQ_SIZE, SQ_SIZE))


if __name__ == "__main__":
    main()   
        
